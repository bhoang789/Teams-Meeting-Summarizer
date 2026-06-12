import pytest
import os
from pathlib import Path
from docx import Document
from src.summary_writer import write_summary_to_docx


@pytest.fixture
def sample_summary():
    """Sample summary data"""
    return {
        'attendees': ['John Smith', 'Jane Doe', 'Bob Johnson'],
        'main_points': [
            {'topic': 'Q3 Planning', 'points': ['Focus on budget', 'Timeline is critical']},
            {'topic': 'Resources', 'points': ['Need 2 engineers', 'Budget allocated']}
        ],
        'action_items': [
            {'task': 'Budget review', 'owner': 'Jane', 'due_date': '2024-06-20'},
            {'task': 'Timeline planning', 'owner': 'Bob', 'due_date': '2024-06-19'}
        ]
    }


class TestSummaryWriter:
    """Test the summary document generation"""
    
    def test_write_summary_to_docx_creates_file(self, tmp_path, sample_summary):
        """
        Given a summary dict and output path,
        When write_summary_to_docx() is called,
        Then a .docx file is created at the specified path
        """
        output_path = tmp_path / "summary.docx"
        
        write_summary_to_docx(sample_summary, str(output_path))
        
        assert output_path.exists()
        assert output_path.suffix == '.docx'
    
    def test_write_summary_contains_attendees(self, tmp_path, sample_summary):
        """
        When write_summary_to_docx() is called,
        Then the document contains all attendees
        """
        output_path = tmp_path / "summary.docx"
        
        write_summary_to_docx(sample_summary, str(output_path))
        
        doc = Document(str(output_path))
        doc_text = '\n'.join([p.text for p in doc.paragraphs])
        
        for attendee in sample_summary['attendees']:
            assert attendee in doc_text
    
    def test_write_summary_contains_main_points(self, tmp_path, sample_summary):
        """
        When write_summary_to_docx() is called,
        Then the document contains all main points
        """
        output_path = tmp_path / "summary.docx"
        
        write_summary_to_docx(sample_summary, str(output_path))
        
        doc = Document(str(output_path))
        doc_text = '\n'.join([p.text for p in doc.paragraphs])
        
        for point_group in sample_summary['main_points']:
            assert point_group['topic'] in doc_text
            for point in point_group['points']:
                assert point in doc_text
    
    def test_write_summary_contains_action_items(self, tmp_path, sample_summary):
        """
        When write_summary_to_docx() is called,
        Then the document contains all action items with owner and due date
        """
        output_path = tmp_path / "summary.docx"
        
        write_summary_to_docx(sample_summary, str(output_path))
        
        doc = Document(str(output_path))
        doc_text = '\n'.join([p.text for p in doc.paragraphs])
        
        for action_item in sample_summary['action_items']:
            assert action_item['task'] in doc_text
            assert action_item['owner'] in doc_text
            assert action_item['due_date'] in doc_text
    
    def test_write_summary_document_structure_has_sections(self, tmp_path, sample_summary):
        """
        When write_summary_to_docx() is called,
        Then the document has clear sections (Attendees, Main Points, Action Items)
        """
        output_path = tmp_path / "summary.docx"
        
        write_summary_to_docx(sample_summary, str(output_path))
        
        doc = Document(str(output_path))
        doc_text = '\n'.join([p.text for p in doc.paragraphs])
        
        # Should have section headers
        assert 'Attendees' in doc_text or 'attendees' in doc_text.lower()
        assert 'Main Points' in doc_text or 'main points' in doc_text.lower()
        assert 'Action Items' in doc_text or 'action items' in doc_text.lower()

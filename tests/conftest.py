import pytest
import os
from docx import Document


@pytest.fixture
def sample_transcript_docx(tmp_path):
    """
    Create a sample Teams meeting transcript .docx file for testing.
    
    Returns:
        str: Path to the created .docx file
    """
    doc = Document()
    
    # Add header
    doc.add_heading('Teams Meeting Transcript', level=1)
    doc.add_paragraph('Date: 2024-06-12')
    doc.add_paragraph('Attendees: John Smith, Jane Doe, Bob Johnson')
    
    # Add some transcript content
    doc.add_heading('Meeting Content', level=2)
    doc.add_paragraph('John: Let\'s discuss the Q3 planning')
    doc.add_paragraph('Jane: We need to focus on three main areas: budget, timeline, and resources')
    doc.add_paragraph('Bob: I agree. The timeline is critical - we have 8 weeks')
    doc.add_paragraph('John: Action items: Budget review by Jane, Timeline planning by Bob')
    
    # Save to temporary file
    file_path = tmp_path / "sample_transcript.docx"
    doc.save(str(file_path))
    
    return str(file_path)

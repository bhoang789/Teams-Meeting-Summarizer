from docx import Document
from datetime import datetime


def write_summary_to_docx(summary_data, output_path):
    """
    Write a structured summary to a Word document.
    
    Args:
        summary_data (dict): Summary data with attendees, main_points, and action_items
        output_path (str): Path where the .docx file should be saved
    
    Returns:
        None
    """
    doc = Document()
    
    # Add title
    doc.add_heading('Meeting Summary', level=1)
    doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
    doc.add_paragraph('')  # Blank line
    
    # Add Attendees section
    doc.add_heading('Attendees', level=2)
    attendees = summary_data.get('attendees', [])
    if attendees:
        for attendee in attendees:
            doc.add_paragraph(attendee, style='List Bullet')
    else:
        doc.add_paragraph('No attendees listed', style='List Bullet')
    
    doc.add_paragraph('')  # Blank line
    
    # Add Main Points section
    doc.add_heading('Main Points', level=2)
    main_points = summary_data.get('main_points', [])
    if main_points:
        for point_group in main_points:
            topic = point_group.get('topic', 'Untitled Topic')
            points = point_group.get('points', [])
            
            doc.add_heading(topic, level=3)
            if points:
                for point in points:
                    doc.add_paragraph(point, style='List Bullet')
            else:
                doc.add_paragraph('(No points listed)')
    else:
        doc.add_paragraph('No main points recorded', style='List Bullet')
    
    doc.add_paragraph('')  # Blank line
    
    # Add Action Items section
    doc.add_heading('Action Items', level=2)
    action_items = summary_data.get('action_items', [])
    if action_items:
        for item in action_items:
            task = item.get('task', 'Untitled Task')
            owner = item.get('owner', 'Unassigned')
            due_date = item.get('due_date', 'No due date')
            
            # Add task as bullet
            p = doc.add_paragraph(f'{task}', style='List Bullet')
            
            # Add owner and due date as sub-items
            doc.add_paragraph(f'Owner: {owner}', style='List Bullet 2')
            doc.add_paragraph(f'Due: {due_date}', style='List Bullet 2')
    else:
        doc.add_paragraph('No action items recorded', style='List Bullet')
    
    # Save the document
    doc.save(output_path)
